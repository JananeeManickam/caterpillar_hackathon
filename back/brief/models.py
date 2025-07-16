# from django.db import models
# from django.core.exceptions import ValidationError
# from os.path import splitext
# import io
# import csv
# import pandas as pd
# import docx
# import PyPDF2

# def validate_file_type(value):
#     allowed_types = ['.pdf', '.pptx', '.docx', '.xlsx', '.csv']
#     ext = splitext(value.name)[1].lower()
#     if ext not in allowed_types:
#         raise ValidationError(f"Unsupported file type: {ext}. Allowed types: {', '.join(allowed_types)}")

# class UploadedFile(models.Model):
#     file = models.FileField(upload_to='uploads/', validators=[validate_file_type])
#     content = models.TextField(blank=True, null=True)
#     uploaded_at = models.DateTimeField(auto_now_add=True)

#     def __str__(self):
#         return self.file.name

#     @property
#     def filename(self):
#         return self.file.name.split('/')[-1]

#     @property
#     def file_extension(self):
#         return splitext(self.file.name)[1].lower()

#     @property
#     def file_size(self):
#         return self.file.size

#     def extract_content(self):
#         """Extract text content from the uploaded file"""
#         if not self.file:
#             return ""
        
#         ext = self.file_extension
        
#         try:
#             self.file.seek(0)
            
#             if ext == '.pdf':
#                 return self._extract_from_pdf()
#             elif ext == '.docx':
#                 return self._extract_from_docx()
#             elif ext == '.xlsx':
#                 return self._extract_from_xlsx()
#             elif ext == '.csv':
#                 return self._extract_from_csv()
            
#         except Exception as e:
#             return f"Error extracting content: {str(e)}"
        
#         return ""

#     def _extract_from_pdf(self):
#         text = ""
#         reader = PyPDF2.PdfReader(self.file)
#         for page in reader.pages:
#             text += page.extract_text() or "" + "\n"
#         return text.strip()

#     def _extract_from_docx(self):
#         doc = docx.Document(self.file)
#         return "\n".join([para.text for para in doc.paragraphs])

#     def _extract_from_xlsx(self):
#         text = ""
#         df = pd.read_excel(self.file, sheet_name=None)
#         for sheet, data in df.items():
#             text += f"Sheet: {sheet}\n"
#             text += data.to_string(index=False, na_rep='') + "\n\n"
#         return text.strip()

#     def _extract_from_csv(self):
#         content = self.file.read().decode('utf-8')
#         csv_reader = csv.reader(io.StringIO(content))
#         return "\n".join([",".join(row) for row in csv_reader])

#     def save(self, *args, **kwargs):
#         # Extract content when saving
#         if self.file and not self.content:
#             self.content = self.extract_content()
#         super().save(*args, **kwargs)

#     class Meta:
#         db_table = 'brief_uploadedfile'
#         ordering = ['-uploaded_at']

from django.db import models
from django.core.exceptions import ValidationError
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from rest_framework.decorators import action
from rest_framework.response import Response
from os.path import splitext
import io
import csv
import pandas as pd
import docx
import PyPDF2

def validate_file_type(value):
    allowed_types = ['.pdf', '.pptx', '.docx', '.xlsx', '.csv']
    ext = splitext(value.name)[1].lower()
    if ext not in allowed_types:
        raise ValidationError(f"Unsupported file type: {ext}. Allowed types: {', '.join(allowed_types)}")

class UploadedFile(models.Model):
    file = models.FileField(upload_to='uploads/', validators=[validate_file_type])
    content = models.TextField(blank=True, null=True)
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.file.name

    @property
    def filename(self):
        return self.file.name.split('/')[-1]

    @property
    def file_extension(self):
        return splitext(self.file.name)[1].lower()

    @property
    def file_size(self):
        return self.file.size

    def extract_content(self):
        """Extract text content from the uploaded file"""
        if not self.file:
            return ""
        
        ext = self.file_extension
        
        try:
            # Reset file pointer to beginning
            self.file.seek(0)
            
            if ext == '.pdf':
                return self._extract_from_pdf()
            elif ext == '.docx':
                return self._extract_from_docx()
            elif ext == '.xlsx':
                return self._extract_from_xlsx()
            elif ext == '.csv':
                return self._extract_from_csv()
            # elif ext == '.pptx':
            #     return self._extract_from_pptx()
            
        except Exception as e:
            return f"Error extracting content: {str(e)}"
        
        return ""

    def _extract_from_pdf(self):
        text = ""
        try:
            reader = PyPDF2.PdfReader(self.file)
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text += f"--- Page {page_num} ---\n"
                    text += page_text + "\n\n"
        except Exception as e:
            text = f"Error reading PDF: {str(e)}"
        return text.strip()

    def _extract_from_docx(self):
        try:
            doc = docx.Document(self.file)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            return "\n".join(paragraphs)
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"

    def _extract_from_xlsx(self):
        text = ""
        try:
            df = pd.read_excel(self.file, sheet_name=None)
            for sheet_name, data in df.items():
                text += f"=== Sheet: {sheet_name} ===\n"
                # Convert to string with better formatting
                text += data.to_string(index=False, na_rep='', max_rows=None, max_cols=None)
                text += "\n\n"
        except Exception as e:
            text = f"Error reading XLSX: {str(e)}"
        return text.strip()

    def _extract_from_csv(self):
        try:
            content = self.file.read().decode('utf-8')
            self.file.seek(0)  # Reset for potential re-reading
            csv_reader = csv.reader(io.StringIO(content))
            rows = []
            for row in csv_reader:
                rows.append(",".join(row))
            return "\n".join(rows)
        except Exception as e:
            return f"Error reading CSV: {str(e)}"

    # def _extract_from_pptx(self):
    #     try:
    #         from pptx import Presentation
    #         prs = Presentation(self.file)
    #         text = ""
    #         for slide_num, slide in enumerate(prs.slides, 1):
    #             text += f"--- Slide {slide_num} ---\n"
    #             for shape in slide.shapes:
    #                 if hasattr(shape, "text") and shape.text.strip():
    #                     text += shape.text + "\n"
    #             text += "\n"
    #         return text.strip()
    #     except ImportError:
    #         return "python-pptx library not installed"
    #     except Exception as e:
    #         return f"Error reading PPTX: {str(e)}"

    def force_extract_content(self):
        """Force re-extraction of content and update the database"""
        self.content = self.extract_content()
        self.save(update_fields=['content'])
        return self.content

    def save(self, *args, **kwargs):
        # Always extract content when saving, even if content exists
        # This ensures content is always up-to-date
        if self.file:
            extracted_content = self.extract_content()
            if extracted_content:
                self.content = extracted_content
        super().save(*args, **kwargs)

    @action(detail=True, methods=['get'])
    def get_content(self, request, pk=None):
        """
        API endpoint to get only the content attribute for a particular file
        Usage: GET /api/uploadedfiles/{id}/get_content/
        """
        try:
            file_obj = get_object_or_404(UploadedFile, pk=pk)
            
            # If content is empty, try to extract it
            if not file_obj.content:
                content = file_obj.extract_content()
                if content:
                    file_obj.content = content
                    file_obj.save(update_fields=['content'])
            
            return Response({
                'id': file_obj.id,
                'filename': file_obj.filename,
                'content': file_obj.content or "No content available",
                'file_size': file_obj.file_size,
                'uploaded_at': file_obj.uploaded_at
            })
        except Exception as e:
            return Response({
                'error': f'Error retrieving content: {str(e)}'
            }, status=500)

    # Alternative manual function if not using DRF
    def get_content_manual(self):
        """Manual function to return content - can be called directly"""
        if not self.content:
            self.content = self.extract_content()
            self.save(update_fields=['content'])
        return self.content

    class Meta:
        db_table = 'brief_uploadedfile'
        ordering = ['-uploaded_at']

# If you want a standalone function outside the model
def get_file_content_by_id(file_id):
    """
    Standalone function to get content by file ID
    Usage: content = get_file_content_by_id(1)
    """
    try:
        file_obj = UploadedFile.objects.get(id=file_id)
        return file_obj.get_content_manual()
    except UploadedFile.DoesNotExist:
        return None
    except Exception as e:
        return f"Error: {str(e)}"